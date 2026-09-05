from django.db import models
from django.contrib.auth import get_user_model
from django.core.validators import FileExtensionValidator
from django.utils import timezone
from django.core.exceptions import ValidationError
from django.conf import settings
import json
import uuid

# Cryptography removed - not needed for basic letter system

User = get_user_model()


# Import Penduduk from references app
try:
    from references.models import Penduduk
    print("Letters app: Using references.models.Penduduk")
except ImportError as e:
    print(f"Letters app: Could not import references.models.Penduduk: {e}")
    # Create fallback Penduduk model
    class Penduduk(models.Model):
        nama = models.CharField(max_length=200)
        nik = models.CharField(max_length=16, unique=True)
        alamat = models.TextField(blank=True)
        telepon = models.CharField(max_length=20, blank=True)
        email = models.EmailField(blank=True)
        created_at = models.DateTimeField(auto_now_add=True)
        updated_at = models.DateTimeField(auto_now=True)
        
        class Meta:
            verbose_name = 'Penduduk'
            verbose_name_plural = 'Penduduk'
        
        def __str__(self):
            return self.nama


class LetterType(models.Model):
    """Model untuk jenis surat"""
    name = models.CharField(max_length=100, verbose_name='Nama Jenis Surat')
    code = models.CharField(max_length=20, unique=True, verbose_name='Kode Surat')
    description = models.TextField(blank=True, verbose_name='Deskripsi')
    template_file = models.FileField(
        upload_to='letter_templates/',
        blank=True,
        null=True,
        validators=[FileExtensionValidator(allowed_extensions=['docx', 'pdf'])],
        verbose_name='File Template'
    )
    required_documents = models.TextField(
        blank=True,
        help_text='Dokumen yang diperlukan (pisahkan dengan koma)',
        verbose_name='Dokumen yang Diperlukan'
    )
    processing_time_days = models.PositiveIntegerField(
        default=3,
        verbose_name='Waktu Proses (Hari)'
    )
    fee_amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0,
        verbose_name='Biaya'
    )
    is_active = models.BooleanField(default=True, verbose_name='Aktif')
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        verbose_name = 'Jenis Surat'
        verbose_name_plural = 'Jenis Surat'
        ordering = ['name']

    def __str__(self):
        return f"{self.code} - {self.name}"


class Letter(models.Model):
    """Model untuk surat"""
    STATUS_CHOICES = [
        ('draft', 'Draft'),
        ('submitted', 'Diajukan'),
        ('in_review', 'Sedang Ditinjau'),
        ('approved', 'Disetujui'),
        ('rejected', 'Ditolak'),
        ('completed', 'Selesai'),
        ('cancelled', 'Dibatalkan'),
    ]

    PRIORITY_CHOICES = [
        ('low', 'Rendah'),
        ('normal', 'Normal'),
        ('high', 'Tinggi'),
        ('urgent', 'Mendesak'),
    ]

    letter_number = models.CharField(
        max_length=50,
        unique=True,
        blank=True,
        verbose_name='Nomor Surat'
    )
    letter_type = models.ForeignKey(
        LetterType,
        on_delete=models.CASCADE,
        verbose_name='Jenis Surat'
    )
    applicant = models.ForeignKey(
        Penduduk,
        on_delete=models.CASCADE,
        verbose_name='Pemohon'
    )
    subject = models.CharField(max_length=200, verbose_name='Perihal')
    content = models.TextField(verbose_name='Isi Surat')
    purpose = models.TextField(verbose_name='Tujuan Penggunaan')
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Status'
    )
    priority = models.CharField(
        max_length=20,
        choices=PRIORITY_CHOICES,
        default='normal',
        verbose_name='Prioritas'
    )
    submission_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Tanggal Pengajuan'
    )
    approval_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Tanggal Persetujuan'
    )
    completion_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Tanggal Selesai'
    )
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        related_name='approved_letters',
        verbose_name='Disetujui Oleh'
    )
    rejection_reason = models.TextField(
        blank=True,
        verbose_name='Alasan Penolakan'
    )
    notes = models.TextField(blank=True, verbose_name='Catatan')
    
    # Template fields
    template = models.ForeignKey(
        'LetterTemplate',
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        verbose_name='Template Surat'
    )
    
    # Digital signature fields removed - simplified system
    
    # Export and sharing
    pdf_file = models.FileField(
        upload_to='letter_pdfs/',
        blank=True,
        null=True,
        verbose_name='File PDF'
    )
    word_file = models.FileField(
        upload_to='letter_docs/',
        blank=True,
        null=True,
        verbose_name='File Word'
    )
    public_url = models.CharField(
        max_length=100,
        blank=True,
        unique=True,
        verbose_name='URL Publik'
    )
    
    # Metadata
    word_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Jumlah Kata'
    )
    estimated_reading_time = models.PositiveIntegerField(
        default=0,
        verbose_name='Estimasi Waktu Baca (detik)'
    )
    language = models.CharField(
        max_length=10,
        default='id',
        verbose_name='Bahasa'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_letters',
        verbose_name='Dibuat Oleh'
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        verbose_name = 'Surat'
        verbose_name_plural = 'Surat'
        ordering = ['-created_at']

    def __str__(self):
        return f"{self.letter_number or 'Draft'} - {self.subject}"

    def save(self, *args, **kwargs):
        # Generate letter number using LetterSettings
        if not self.letter_number and self.status != 'draft':
            try:
                letter_settings = LetterSettings.objects.filter(is_active=True).first()
                if letter_settings:
                    self.letter_number = letter_settings.get_next_letter_number(self.letter_type.code)
                else:
                    # Fallback to old method
                    year = timezone.now().year
                    month = timezone.now().month
                    count = Letter.objects.filter(
                        created_at__year=year,
                        created_at__month=month
                    ).count() + 1
                    self.letter_number = f"{self.letter_type.code}/{count:03d}/{month:02d}/{year}"
            except:
                # Fallback to old method if LetterSettings not available
                year = timezone.now().year
                month = timezone.now().month
                count = Letter.objects.filter(
                    created_at__year=year,
                    created_at__month=month
                ).count() + 1
                self.letter_number = f"{self.letter_type.code}/{count:03d}/{month:02d}/{year}"
        
        # Generate public URL if not exists
        if not self.public_url:
            self.public_url = str(uuid.uuid4())[:8]
        
        # Calculate word count and reading time
        if self.content:
            self.word_count = len(self.content.split())
            # Average reading speed: 200 words per minute
            self.estimated_reading_time = max(1, (self.word_count * 60) // 200)
        
        super().save(*args, **kwargs)
        
        # QR code generation removed
    
    # QR code generation method removed
    
    def generate_pdf(self):
        """Generate PDF version of the letter for printing"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import inch
            from io import BytesIO
            from django.core.files import File
            
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=A4)
            width, height = A4
            
            # Get letter settings for header
            letter_settings = LetterSettings.objects.filter(is_active=True).first()
            
            # Add content to PDF
            y_position = height - 100
            
            if letter_settings:
                # Header with village name and address
                p.setFont("Helvetica-Bold", 16)
                p.drawCentredText(width/2, y_position, letter_settings.village_name)
                y_position -= 30
                p.setFont("Helvetica", 12)
                p.drawCentredText(width/2, y_position, letter_settings.village_address)
                y_position -= 50
            
            # Letter number and date
            p.setFont("Helvetica", 12)
            p.drawString(100, y_position, f"Nomor: {self.letter_number or 'Draft'}")
            y_position -= 25
            p.drawString(100, y_position, f"Tanggal: {self.created_at.strftime('%d %B %Y')}")
            y_position -= 50
            
            # Subject
            p.setFont("Helvetica-Bold", 14)
            p.drawString(100, y_position, f"Perihal: {self.subject}")
            y_position -= 50
            
            # Content with proper text wrapping
            p.setFont("Helvetica", 12)
            content_lines = self._wrap_text(self.content, 80)
            for line in content_lines:
                if y_position < 150:  # Start new page if needed
                    p.showPage()
                    y_position = height - 100
                p.drawString(100, y_position, line)
                y_position -= 20
            
            # Signature area
            if letter_settings and y_position > 200:
                y_position -= 50
                p.drawString(width - 200, y_position, f"{letter_settings.village_name}")
                y_position -= 30
                p.drawString(width - 200, y_position, f"Kepala Desa")
                y_position -= 80
                p.drawString(width - 200, y_position, f"{letter_settings.head_of_village_name}")
            
            p.save()
            buffer.seek(0)
            
            # Save PDF file
            filename = f"letter_{self.public_url}.pdf"
            self.pdf_file.save(filename, File(buffer), save=False)
            self.save(update_fields=['pdf_file'])
            
            return True
            
        except ImportError:
            # ReportLab not installed - this is optional
            return False
        except Exception as e:
            return False
    
    def _wrap_text(self, text, max_length):
        """Wrap text to fit within specified length"""
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            if len(current_line + " " + word) <= max_length:
                if current_line:
                    current_line += " " + word
                else:
                    current_line = word
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def generate_word_document(self):
        """Generate Word document version of the letter for printing"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            from django.core.files import File
            
            # Create new document
            doc = Document()
            
            # Get letter settings
            letter_settings = LetterSettings.objects.filter(is_active=True).first()
            
            # Add header
            if letter_settings:
                header = doc.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header.add_run(letter_settings.village_name)
                header_run.bold = True
                header_run.font.size = Inches(0.2)
                
                address = doc.add_paragraph()
                address.alignment = WD_ALIGN_PARAGRAPH.CENTER
                address.add_run(letter_settings.village_address)
            
            # Add letter details
            doc.add_paragraph()  # Empty line
            
            details = doc.add_paragraph()
            details.add_run(f"Nomor: {self.letter_number or 'Draft'}")
            details.add_run(f"\nTanggal: {self.created_at.strftime('%d %B %Y')}")
            
            doc.add_paragraph()  # Empty line
            
            # Add subject
            subject = doc.add_paragraph()
            subject_run = subject.add_run(f"Perihal: {self.subject}")
            subject_run.bold = True
            
            doc.add_paragraph()  # Empty line
            
            # Add content
            content_para = doc.add_paragraph()
            content_para.add_run(self.content)
            
            # Add signature area
            if letter_settings:
                doc.add_paragraph()  # Empty line
                doc.add_paragraph()  # Empty line
                
                signature = doc.add_paragraph()
                signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                signature.add_run(f"{letter_settings.village_name}")
                signature.add_run(f"\nKepala Desa")
                signature.add_run(f"\n\n{letter_settings.head_of_village_name}")
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Save Word file
            filename = f"letter_{self.public_url}.docx"
            self.word_file.save(filename, File(buffer), save=False)
            self.save(update_fields=['word_file'])
            
            return True
            
        except ImportError:
            # python-docx not installed
            return False
        except Exception as e:
            return False
    
    def get_verification_url(self):
        """Get public verification URL"""
        return f"https://pulosarok.desa.id/verify/{self.public_url}"
    
    def print_letter(self, format_type='pdf'):
        """Generate and return file for printing"""
        if format_type.lower() == 'pdf':
            if not self.pdf_file:
                self.generate_pdf()
            return self.pdf_file
        elif format_type.lower() == 'word':
            if not self.word_file:
                self.generate_word_document()
            return self.word_file
        else:
            return None
    
    def get_print_url(self, format_type='pdf'):
        """Get URL for printing the letter"""
        if format_type.lower() == 'pdf':
            if not self.pdf_file:
                self.generate_pdf()
            return self.pdf_file.url if self.pdf_file else None
        elif format_type.lower() == 'word':
            if not self.word_file:
                self.generate_word_document()
            return self.word_file.url if self.word_file else None
        else:
            return None
    
    def calculate_similarity_score(self, other_content):
        """Calculate similarity score with another content"""
        try:
            from difflib import SequenceMatcher
            return SequenceMatcher(None, self.content, other_content).ratio()
        except:
            return 0.0
    
    # AI validation methods removed


class LetterRecipient(models.Model):
    """Model untuk penerima surat"""
    RECIPIENT_TYPE_CHOICES = [
        ('internal', 'Internal'),
        ('external', 'Eksternal'),
        ('government', 'Instansi Pemerintah'),
        ('private', 'Swasta'),
        ('individual', 'Perorangan'),
    ]

    letter = models.ForeignKey(
        Letter,
        on_delete=models.CASCADE,
        related_name='recipients',
        verbose_name='Surat'
    )
    recipient_type = models.CharField(
        max_length=20,
        choices=RECIPIENT_TYPE_CHOICES,
        verbose_name='Jenis Penerima'
    )
    name = models.CharField(max_length=200, verbose_name='Nama Penerima')
    position = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Jabatan'
    )
    organization = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Organisasi/Instansi'
    )
    address = models.TextField(verbose_name='Alamat')
    phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Telepon'
    )
    email = models.EmailField(blank=True, verbose_name='Email')
    is_primary = models.BooleanField(
        default=False,
        verbose_name='Penerima Utama'
    )
    delivery_method = models.CharField(
        max_length=20,
        choices=[
            ('hand_delivery', 'Antar Langsung'),
            ('post', 'Pos'),
            ('email', 'Email'),
            ('fax', 'Fax'),
        ],
        default='hand_delivery',
        verbose_name='Metode Pengiriman'
    )
    delivery_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Tanggal Pengiriman'
    )
    received_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Tanggal Diterima'
    )
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Penerima Surat'
        verbose_name_plural = 'Penerima Surat'
        ordering = ['-is_primary', 'name']

    def __str__(self):
        return f"{self.name} - {self.letter.subject}"


class LetterAttachment(models.Model):
    """Model untuk lampiran surat"""
    ATTACHMENT_TYPE_CHOICES = [
        ('supporting_document', 'Dokumen Pendukung'),
        ('identity_card', 'Kartu Identitas'),
        ('certificate', 'Sertifikat'),
        ('photo', 'Foto'),
        ('other', 'Lainnya'),
    ]

    letter = models.ForeignKey(
        Letter,
        on_delete=models.CASCADE,
        related_name='attachments',
        verbose_name='Surat'
    )
    attachment_type = models.CharField(
        max_length=30,
        choices=ATTACHMENT_TYPE_CHOICES,
        verbose_name='Jenis Lampiran'
    )
    title = models.CharField(max_length=200, verbose_name='Judul Lampiran')
    description = models.TextField(blank=True, verbose_name='Deskripsi')
    file = models.FileField(
        upload_to='letter_attachments/',
        validators=[FileExtensionValidator(allowed_extensions=['pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx'])],
        verbose_name='File'
    )
    file_size = models.PositiveIntegerField(
        blank=True,
        null=True,
        verbose_name='Ukuran File (bytes)'
    )
    is_required = models.BooleanField(
        default=False,
        verbose_name='Wajib'
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Diunggah Oleh'
    )
    uploaded_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name = 'Lampiran Surat'
        verbose_name_plural = 'Lampiran Surat'
        ordering = ['-is_required', 'title']

    def __str__(self):
        return f"{self.title} - {self.letter.subject}"

    def save(self, *args, **kwargs):
        if self.file:
            self.file_size = self.file.size
        super().save(*args, **kwargs)


class LetterTracking(models.Model):
    """Model untuk tracking status surat"""
    ACTION_CHOICES = [
        ('created', 'Dibuat'),
        ('submitted', 'Diajukan'),
        ('reviewed', 'Ditinjau'),
        ('approved', 'Disetujui'),
        ('rejected', 'Ditolak'),
        ('completed', 'Diselesaikan'),
        ('cancelled', 'Dibatalkan'),
        ('sent', 'Dikirim'),
        ('received', 'Diterima'),
        ('returned', 'Dikembalikan'),
    ]

    letter = models.ForeignKey(
        Letter,
        on_delete=models.CASCADE,
        related_name='tracking_history',
        verbose_name='Surat'
    )
    action = models.CharField(
        max_length=20,
        choices=ACTION_CHOICES,
        verbose_name='Aksi'
    )
    description = models.TextField(verbose_name='Deskripsi')
    performed_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Dilakukan Oleh'
    )
    performed_at = models.DateTimeField(auto_now_add=True)
    location = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Lokasi'
    )
    notes = models.TextField(blank=True, verbose_name='Catatan')
    ip_address = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name='Alamat IP'
    )

    class Meta:
        verbose_name = 'Tracking Surat'
        verbose_name_plural = 'Tracking Surat'
        ordering = ['-performed_at']

    def __str__(self):
        return f"{self.letter.subject} - {self.get_action_display()}"


# APIKeySettings model removed - AI functionality disabled


class LetterSettings(models.Model):
    """Model untuk konfigurasi surat (kop surat, nomor, kepala desa)"""
    SIGNATURE_TYPE_CHOICES = [
        ('digital', 'Digital'),
        ('image', 'Gambar'),
        ('text', 'Teks'),
    ]

    village_name = models.CharField(
        max_length=100,
        default='Desa Pulosarok',
        verbose_name='Nama Desa'
    )
    village_address = models.TextField(
        default='Kecamatan Pulosarok, Kabupaten Indramayu',
        verbose_name='Alamat Desa'
    )
    village_phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Telepon Desa'
    )
    village_email = models.EmailField(
        blank=True,
        verbose_name='Email Desa'
    )
    village_website = models.URLField(
        blank=True,
        verbose_name='Website Desa'
    )
    village_logo = models.ImageField(
        upload_to='village_settings/',
        blank=True,
        null=True,
        verbose_name='Logo Desa'
    )
    
    # Kepala Desa
    head_of_village_name = models.CharField(
        max_length=100,
        verbose_name='Nama Kepala Desa'
    )
    head_of_village_nip = models.CharField(
        max_length=30,
        blank=True,
        verbose_name='NIP Kepala Desa'
    )
    head_of_village_signature_type = models.CharField(
        max_length=20,
        choices=SIGNATURE_TYPE_CHOICES,
        default='digital',
        verbose_name='Jenis Tanda Tangan'
    )
    head_of_village_signature_image = models.ImageField(
        upload_to='signatures/',
        blank=True,
        null=True,
        verbose_name='Gambar Tanda Tangan'
    )
    head_of_village_signature_text = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Teks Tanda Tangan'
    )
    
    # Sekretaris Desa
    secretary_name = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Nama Sekretaris Desa'
    )
    secretary_nip = models.CharField(
        max_length=30,
        blank=True,
        verbose_name='NIP Sekretaris'
    )
    secretary_signature_type = models.CharField(
        max_length=20,
        choices=SIGNATURE_TYPE_CHOICES,
        default='digital',
        verbose_name='Jenis Tanda Tangan Sekretaris'
    )
    secretary_signature_image = models.ImageField(
        upload_to='signatures/',
        blank=True,
        null=True,
        verbose_name='Gambar Tanda Tangan Sekretaris'
    )
    
    # Pengaturan Nomor Surat
    letter_number_format = models.CharField(
        max_length=100,
        default='{code}/{number:03d}/{month:02d}/{year}',
        verbose_name='Format Nomor Surat',
        help_text='Gunakan {code}, {number}, {month}, {year} sebagai placeholder'
    )
    current_year_counter = models.PositiveIntegerField(
        default=0,
        verbose_name='Counter Tahun Ini'
    )
    reset_counter_yearly = models.BooleanField(
        default=True,
        verbose_name='Reset Counter Setiap Tahun'
    )
    
    # AI Settings removed - simplified system
    
    # Digital Signature Settings
    enable_digital_signature = models.BooleanField(
        default=True,
        verbose_name='Aktifkan Tanda Tangan Digital'
    )
    signature_certificate = models.FileField(
        upload_to='certificates/',
        blank=True,
        null=True,
        verbose_name='Sertifikat Digital'
    )
    
    is_active = models.BooleanField(default=True, verbose_name='Aktif')
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Dibuat Oleh'
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        verbose_name = 'Pengaturan Surat'
        verbose_name_plural = 'Pengaturan Surat'
        ordering = ['-is_active', '-created_at']

    def __str__(self):
        return f"Pengaturan {self.village_name}"
    
    def save(self, *args, **kwargs):
        # Only one active settings at a time
        if self.is_active:
            LetterSettings.objects.filter(is_active=True).update(is_active=False)
        super().save(*args, **kwargs)

    def get_next_letter_number(self, letter_type_code):
        """Generate next letter number"""
        now = timezone.now()
        
        if self.reset_counter_yearly:
            # Reset counter if new year
            current_year = now.year
            if not hasattr(self, '_last_year') or self._last_year != current_year:
                self.current_year_counter = 0
                self._last_year = current_year
        
        self.current_year_counter += 1
        self.save()
        
        return self.letter_number_format.format(
            code=letter_type_code,
            number=self.current_year_counter,
            month=now.month,
            year=now.year
        )


class LetterTemplate(models.Model):
    """Model untuk template surat"""
    TEMPLATE_TYPE_CHOICES = [
        ('official', 'Surat Resmi'),
        ('certificate', 'Surat Keterangan'),
        ('recommendation', 'Surat Rekomendasi'),
        ('invitation', 'Surat Undangan'),
        ('notification', 'Surat Pemberitahuan'),
        ('custom', 'Template Kustom'),
    ]

    name = models.CharField(max_length=100, verbose_name='Nama Template')
    template_type = models.CharField(
        max_length=20,
        choices=TEMPLATE_TYPE_CHOICES,
        verbose_name='Jenis Template'
    )
    description = models.TextField(blank=True, verbose_name='Deskripsi')
    content_template = models.TextField(
        verbose_name='Template Konten',
        help_text='Gunakan {{variable}} untuk placeholder'
    )
    variables = models.TextField(
        default='{}',
        verbose_name='Variabel Template',
        help_text='JSON object dengan daftar variabel yang tersedia'
    )
    css_styles = models.TextField(
        blank=True,
        verbose_name='CSS Styles',
        help_text='CSS untuk styling template'
    )
    header_template = models.TextField(
        blank=True,
        verbose_name='Template Header'
    )
    footer_template = models.TextField(
        blank=True,
        verbose_name='Template Footer'
    )
    is_default = models.BooleanField(
        default=False,
        verbose_name='Template Default'
    )
    is_active = models.BooleanField(default=True, verbose_name='Aktif')
    usage_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Jumlah Penggunaan'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Dibuat Oleh'
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        verbose_name = 'Template Surat'
        verbose_name_plural = 'Template Surat'
        ordering = ['-is_default', '-usage_count', 'name']

    def __str__(self):
        return f"{self.name} ({self.get_template_type_display()})"

    def increment_usage(self):
        """Increment usage counter"""
        self.usage_count += 1
        self.save()

    def render_content(self, context):
        """Render template with context variables"""
        import re
        content = self.content_template
        
        # Replace {{variable}} with context values
        for key, value in context.items():
            pattern = r'\{\{\s*' + re.escape(key) + r'\s*\}\}'
            content = re.sub(pattern, str(value), content)
        
        return content


# LetterAIValidation model removed - AI functionality disabled


# LetterDigitalSignature model removed - simplified system
